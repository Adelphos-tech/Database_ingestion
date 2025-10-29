import os
import json
import logging
from collections import deque
from urllib.parse import urlparse, urljoin, urldefrag
from flask import Flask, request, jsonify
from flask_cors import CORS
from pinecone import Pinecone, ServerlessSpec
import google.generativeai as genai
from google.api_core.exceptions import NotFound, GoogleAPICallError
import hashlib
from datetime import datetime
import PyPDF2
import docx
import pandas as pd
import requests
from bs4 import BeautifulSoup
from io import BytesIO
import traceback
from dotenv import load_dotenv
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import Document as LlamaDocument

try:
    import cloudscraper
except ImportError:
    cloudscraper = None

try:
    from playwright.sync_api import sync_playwright
except ImportError:
    sync_playwright = None

try:
    from playwright_stealth import stealth_sync
except ImportError:
    stealth_sync = None

try:
    import undetected_chromedriver as uc
except ImportError:
    uc = None

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*", "methods": ["GET", "POST", "DELETE", "OPTIONS"], "allow_headers": "*"}})

# Configuration - Set these in environment variables
PINECONE_API_KEY = os.getenv('PINECONE_API_KEY')
PINECONE_ENVIRONMENT = os.getenv('PINECONE_ENVIRONMENT', 'us-east-1')
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
DEFAULT_CHAT_MODEL = 'models/gemini-2.0-flash-exp'
CONFIGURED_CHAT_MODEL = os.getenv('GEMINI_CHAT_MODEL', DEFAULT_CHAT_MODEL)
ENABLE_PLAYWRIGHT_CRAWL = os.getenv('ENABLE_PLAYWRIGHT_CRAWL', 'false').lower() == 'true'

# Custom captcha bypass (no 3rd party APIs)
from captcha_bypass import CustomCaptchaBypass

# Proxy configuration (supports http, https, socks5)
PROXY_URL = os.getenv('PROXY_URL', '')  # Format: http://user:pass@host:port or socks5://host:port
PROXY_ROTATION_ENABLED = os.getenv('PROXY_ROTATION_ENABLED', 'false').lower() == 'true'
PROXY_LIST = os.getenv('PROXY_LIST', '').split(',') if os.getenv('PROXY_LIST') else []

# Initialize clients
pc = Pinecone(api_key=PINECONE_API_KEY)

# Initialize Google Gemini for embeddings
print("Configuring Google Gemini embeddings...")
if not GEMINI_API_KEY:
    print("⚠️ GEMINI_API_KEY not set. Embedding requests will fail until it is provided.")
else:
    genai.configure(api_key=GEMINI_API_KEY)
    print(f"✅ Google Gemini chat model configured (requested): {CONFIGURED_CHAT_MODEL}")
    print("✅ Google Gemini embeddings configured!")
print("✅ Using text-embedding-004 model (768 dimensions)")
print("✅ Matches n8n workflow embedding model")

# Initialize chat model (lazy load fallback inside routes)
chat_model = None
active_chat_model_name = None


def initialize_chat_model(force_fallback=False):
    """Initialize Gemini chat model with graceful fallback when needed."""
    global active_chat_model_name
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY is required for chat functionality.")

    requested_model = CONFIGURED_CHAT_MODEL
    candidate_models = [requested_model]
    if force_fallback or requested_model != DEFAULT_CHAT_MODEL:
        candidate_models.append(DEFAULT_CHAT_MODEL)

    last_error = None
    for model_name in candidate_models:
        try:
            model = genai.GenerativeModel(model_name)
            active_chat_model_name = model_name
            print(f"✅ Gemini chat model active: {model_name}")
            return model
        except NotFound as not_found_error:
            print(f"⚠️ Gemini model '{model_name}' not found. Trying fallback...")
            last_error = not_found_error
        except GoogleAPICallError as api_error:
            print(f"⚠️ Gemini API error initialising model '{model_name}': {api_error}")
            last_error = api_error

    raise RuntimeError(f"Unable to initialise Gemini chat model: {last_error}")


# Index configuration
DEFAULT_INDEX_NAME = "document-knowledge-base"
EMBEDDING_DIMENSION = 768  # Google Gemini text-embedding-004 produces 768-dimensional embeddings

# Chunking configuration (using LlamaIndex's SentenceSplitter - Recursive Character Text Splitter)
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 300  # Increased overlap for better context retrieval


def _get_index_description(index_name):
    """Return Pinecone index description or None if not found."""
    try:
        description = pc.describe_index(index_name)
        if isinstance(description, dict):
            return description
        return description
    except Exception:
        return None


def _extract_dimension(index_description):
    """Best-effort extraction of dimension value from Pinecone index description."""
    if not index_description:
        return None
    for attr in ("dimension", "dimensions"):
        if hasattr(index_description, attr):
            value = getattr(index_description, attr)
            if value:
                return value
    if isinstance(index_description, dict):
        for key in ("dimension", "dimensions"):
            if key in index_description and index_description[key]:
                return index_description[key]
    return None


def get_index_dimension(index_name, fallback=EMBEDDING_DIMENSION):
    """Return the dimension for the given index, falling back to configured embedding dimension."""
    description = _get_index_description(index_name)
    existing_dimension = _extract_dimension(description)
    return existing_dimension or fallback


def get_or_create_index(index_name=None):
    """Get existing index or create new one"""
    try:
        if index_name is None:
            index_name = DEFAULT_INDEX_NAME

        existing_indexes = {index.name: index for index in pc.list_indexes()}
        index_info = existing_indexes.get(index_name)

        if index_info:
            existing_dimension = getattr(index_info, 'dimension', None)
            if existing_dimension is None:
                description = _get_index_description(index_name)
                existing_dimension = _extract_dimension(description)
            if existing_dimension and existing_dimension != EMBEDDING_DIMENSION:
                raise ValueError(
                    f"Pinecone index '{index_name}' is dimension {existing_dimension}, "
                    f"but this service requires {EMBEDDING_DIMENSION}. "
                    "Please recreate the index or choose another one that matches the embedding model."
                )
        else:
            pc.create_index(
                name=index_name,
                dimension=EMBEDDING_DIMENSION,
                metric='cosine',
                spec=ServerlessSpec(
                    cloud='aws',
                    region=PINECONE_ENVIRONMENT
                )
        )

        return pc.Index(index_name)
    except Exception as e:
        print(f"Error creating/getting index: {e}")
        raise


def extract_text_from_pdf(file_bytes):
    """Extract text from PDF file"""
    pdf_file = BytesIO(file_bytes)
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n\n"
    return text


def extract_text_from_docx(file_bytes):
    """Extract text from DOCX file"""
    doc_file = BytesIO(file_bytes)
    doc = docx.Document(doc_file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text


def extract_text_from_excel(file_bytes, filename):
    """Extract text from Excel file"""
    excel_file = BytesIO(file_bytes)
    
    if filename.endswith('.csv'):
        df = pd.read_csv(excel_file)
        return df.to_string()
    else:
        xls = pd.ExcelFile(excel_file)
        text = ""
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name)
            text += f"\n\n=== Sheet: {sheet_name} ===\n"
            text += df.to_string()
        return text


def extract_text(file_bytes, filename):
    """Route to appropriate text extraction based on file type (Default Data Loader functionality)"""
    extension = filename.lower().split('.')[-1]
    
    if extension == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif extension in ['docx', 'doc']:
        return extract_text_from_docx(file_bytes)
    elif extension in ['xlsx', 'xls', 'csv']:
        return extract_text_from_excel(file_bytes, filename)
    elif extension == 'txt':
        return file_bytes.decode('utf-8')
    else:
        raise ValueError(f"Unsupported file type: {extension}")


def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """Split text into overlapping chunks using LlamaIndex's Recursive Character Text Splitter"""
    # Create LlamaIndex Document
    document = LlamaDocument(text=text)
    
    # Initialize SentenceSplitter (LlamaIndex's recursive character text splitter)
    text_splitter = SentenceSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separator=" ",  # Default separator
    )
    
    # Split the document into nodes
    nodes = text_splitter.get_nodes_from_documents([document])
    
    # Convert nodes to our chunk format
    chunks = []
    for idx, node in enumerate(nodes):
        chunks.append({
            'text': node.text,
            'start_position': node.start_char_idx if hasattr(node, 'start_char_idx') else 0,
            'end_position': node.end_char_idx if hasattr(node, 'end_char_idx') else len(node.text),
            'chunk_index': idx
        })

    return chunks


def ingest_text_payload(text, source_name, project, index_name, extra_metadata=None):
    """Chunk, embed, and store text in Pinecone."""
    if not text or not text.strip():
        raise ValueError("No text content to ingest.")

    extra_metadata = extra_metadata.copy() if extra_metadata else {}
    chunks = chunk_text(text)

    if not chunks:
        raise ValueError("Unable to generate chunks from the supplied text.")

    chunk_texts = [chunk['text'] for chunk in chunks]
    embeddings = generate_embeddings(chunk_texts)

    index = get_or_create_index(index_name)
    doc_id = generate_document_id(source_name, project)

    file_size = extra_metadata.get('file_size', len(text))
    source_label = extra_metadata.get('source', source_name)

    base_metadata = {
        'document_id': doc_id,
        'filename': source_name,
        'project': project,
        'total_chunks': len(chunks),
        'upload_date': datetime.now().isoformat(),
        'file_size': file_size,
        'source': source_label
    }

    vectors = []
    for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        metadata = {
            **base_metadata,
            'chunk_index': i,
            'chunk_start': chunk['start_position'],
            'chunk_end': chunk['end_position'],
            'text': chunk['text']
        }
        for key, value in extra_metadata.items():
            if key not in metadata:
                metadata[key] = value

        vector_id = f"{doc_id}_chunk_{i}"
        vectors.append({
            'id': vector_id,
            'values': embedding,
            'metadata': metadata
        })

    batch_size = 100
    for i in range(0, len(vectors), batch_size):
        batch = vectors[i:i + batch_size]
        index.upsert(vectors=batch, namespace=project)

    return {
        'document_id': doc_id,
        'filename': source_name,
        'project': project,
        'chunks_created': len(chunks),
        'total_characters': len(text)
    }


def _normalize_embedding_response(response):
    """Normalize Gemini embedding response into a flat list of floats."""
    embedding = getattr(response, 'embedding', response)
    if embedding is None and isinstance(response, dict):
        embedding = response.get('embedding')
    # Unwrap nested dict structures returned by the SDK
    visited = set()
    while isinstance(embedding, dict):
        dict_id = id(embedding)
        if dict_id in visited:
            break
        visited.add(dict_id)
        if 'values' in embedding:
            embedding = embedding['values']
        elif 'embedding' in embedding:
            embedding = embedding['embedding']
        elif 'data' in embedding:
            embedding = embedding['data']
        elif 'vector' in embedding:
            embedding = embedding['vector']
        else:
            break
    # Handle SDK objects that expose .values attribute/property
    if hasattr(embedding, 'values'):
        values_attr = embedding.values
        if callable(values_attr):
            embedding = values_attr()
        else:
            embedding = values_attr
    if isinstance(embedding, (tuple, set)):
        embedding = list(embedding)
    if not isinstance(embedding, list):
        raise ValueError(f"Unexpected embedding response format: {type(response)}")
    if len(embedding) != EMBEDDING_DIMENSION:
        raise ValueError(f"Embedding dimension mismatch. Expected {EMBEDDING_DIMENSION}, got {len(embedding)}")
    return embedding


def generate_embeddings(texts):
    """Generate embeddings using Google Gemini API (matches n8n workflow)"""
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY is not configured. Set the environment variable to enable embeddings.")
    
    embeddings = []
    for idx, text in enumerate(texts):
        try:
            result = genai.embed_content(
                model="models/text-embedding-004",
                content=text,
                task_type="retrieval_document"
            )
            embeddings.append(_normalize_embedding_response(result))
        except Exception as e:
            print(f"Error generating embedding for chunk {idx}: {e}")
            raise
    
    return embeddings


def generate_document_id(filename, project):
    """Generate unique document ID"""
    timestamp = datetime.now().isoformat()
    unique_string = f"{filename}_{project}_{timestamp}"
    return hashlib.md5(unique_string.encode()).hexdigest()


HEADLESS_HEADERS = [
    {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.9',
        'Accept-Encoding': 'gzip, deflate, br',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1',
        'Sec-Fetch-Dest': 'document',
        'Sec-Fetch-Mode': 'navigate',
        'Sec-Fetch-Site': 'none',
        'Sec-Fetch-User': '?1',
        'Cache-Control': 'max-age=0'
    },
    {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.9',
        'Accept-Encoding': 'gzip, deflate, br',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1'
    },
    {
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.5',
        'Accept-Encoding': 'gzip, deflate, br',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1'
    }
]


# Old CaptchaSolver class removed - now using CustomCaptchaBypass
# No 3rd party APIs needed!
    
    @staticmethod
    def _solve_anticaptcha_v3(site_key, page_url, action, api_key):
        """Solve reCAPTCHA v3 using AntiCaptcha service"""
        try:
            import time
            submit_url = 'https://api.anti-captcha.com/createTask'
            payload = {
                'clientKey': api_key,
                'task': {
                    'type': 'RecaptchaV3TaskProxyless',
                    'websiteURL': page_url,
                    'websiteKey': site_key,
                    'minScore': 0.3,
                    'pageAction': action
                }
            }
            response = requests.post(submit_url, json=payload, timeout=30)
            result = response.json()
            
            if result.get('errorId') != 0:
                return None
            
            task_id = result.get('taskId')
            result_url = 'https://api.anti-captcha.com/getTaskResult'
            
            for _ in range(60):
                time.sleep(2)
                payload = {'clientKey': api_key, 'taskId': task_id}
                response = requests.post(result_url, json=payload, timeout=30)
                result = response.json()
                
                if result.get('status') == 'ready':
                    return result.get('solution', {}).get('gRecaptchaResponse')
                elif result.get('status') != 'processing':
                    return None
            
            return None
        except Exception as e:
            logging.error(f"AntiCaptcha v3 error: {e}")
            return None
    
    @staticmethod
    def _solve_anticaptcha_hcaptcha(site_key, page_url, api_key):
        """Solve hCaptcha using AntiCaptcha service"""
        try:
            import time
            submit_url = 'https://api.anti-captcha.com/createTask'
            payload = {
                'clientKey': api_key,
                'task': {
                    'type': 'HCaptchaTaskProxyless',
                    'websiteURL': page_url,
                    'websiteKey': site_key
                }
            }
            response = requests.post(submit_url, json=payload, timeout=30)
            result = response.json()
            
            if result.get('errorId') != 0:
                return None
            
            task_id = result.get('taskId')
            result_url = 'https://api.anti-captcha.com/getTaskResult'
            
            for _ in range(60):
                time.sleep(2)
                payload = {'clientKey': api_key, 'taskId': task_id}
                response = requests.post(result_url, json=payload, timeout=30)
                result = response.json()
                
                if result.get('status') == 'ready':
                    return result.get('solution', {}).get('gRecaptchaResponse')
                elif result.get('status') != 'processing':
                    return None
            
            return None
        except Exception as e:
            logging.error(f"AntiCaptcha hCaptcha error: {e}")
            return None
    
    @staticmethod
    def _solve_capsolver_v2(site_key, page_url, api_key):
        """Solve reCAPTCHA v2 using CapSolver service"""
        try:
            import time
            submit_url = 'https://api.capsolver.com/createTask'
            payload = {
                'clientKey': api_key,
                'task': {
                    'type': 'ReCaptchaV2TaskProxyLess',
                    'websiteURL': page_url,
                    'websiteKey': site_key
                }
            }
            response = requests.post(submit_url, json=payload, timeout=30)
            result = response.json()
            
            if result.get('errorId') != 0:
                return None
            
            task_id = result.get('taskId')
            result_url = 'https://api.capsolver.com/getTaskResult'
            
            for _ in range(60):
                time.sleep(2)
                payload = {'clientKey': api_key, 'taskId': task_id}
                response = requests.post(result_url, json=payload, timeout=30)
                result = response.json()
                
                if result.get('status') == 'ready':
                    return result.get('solution', {}).get('gRecaptchaResponse')
                elif result.get('status') != 'processing':
                    return None
            
            return None
        except Exception as e:
            logging.error(f"CapSolver error: {e}")
            return None
    
    @staticmethod
    def _solve_capsolver_v3(site_key, page_url, action, api_key):
        """Solve reCAPTCHA v3 using CapSolver service"""
        try:
            import time
            submit_url = 'https://api.capsolver.com/createTask'
            payload = {
                'clientKey': api_key,
                'task': {
                    'type': 'ReCaptchaV3TaskProxyLess',
                    'websiteURL': page_url,
                    'websiteKey': site_key,
                    'pageAction': action
                }
            }
            response = requests.post(submit_url, json=payload, timeout=30)
            result = response.json()
            
            if result.get('errorId') != 0:
                return None
            
            task_id = result.get('taskId')
            result_url = 'https://api.capsolver.com/getTaskResult'
            
            for _ in range(60):
                time.sleep(2)
                payload = {'clientKey': api_key, 'taskId': task_id}
                response = requests.post(result_url, json=payload, timeout=30)
                result = response.json()
                
                if result.get('status') == 'ready':
                    return result.get('solution', {}).get('gRecaptchaResponse')
                elif result.get('status') != 'processing':
                    return None
            
            return None
        except Exception as e:
            logging.error(f"CapSolver v3 error: {e}")
            return None
    
    @staticmethod
    def _solve_capsolver_hcaptcha(site_key, page_url, api_key):
        """Solve hCaptcha using CapSolver service"""
        try:
            import time
            submit_url = 'https://api.capsolver.com/createTask'
            payload = {
                'clientKey': api_key,
                'task': {
                    'type': 'HCaptchaTaskProxyLess',
                    'websiteURL': page_url,
                    'websiteKey': site_key
                }
            }
            response = requests.post(submit_url, json=payload, timeout=30)
            result = response.json()
            
            if result.get('errorId') != 0:
                return None
            
            task_id = result.get('taskId')
            result_url = 'https://api.capsolver.com/getTaskResult'
            
            for _ in range(60):
                time.sleep(2)
                payload = {'clientKey': api_key, 'taskId': task_id}
                response = requests.post(result_url, json=payload, timeout=30)
                result = response.json()
                
                if result.get('status') == 'ready':
                    return result.get('solution', {}).get('gRecaptchaResponse')
                elif result.get('status') != 'processing':
                    return None
            
            return None
        except Exception as e:
            logging.error(f"CapSolver hCaptcha error: {e}")
            return None


def get_random_proxy():
    """Get a random proxy from the proxy list"""
    if PROXY_ROTATION_ENABLED and PROXY_LIST:
        import random
        return random.choice(PROXY_LIST).strip()
    return PROXY_URL if PROXY_URL else None


def fetch_with_requests(url, timeout):
    """Fetch with requests library supporting proxies"""
    proxy = get_random_proxy()
    proxies = {'http': proxy, 'https': proxy} if proxy else None
    
    for headers in HEADLESS_HEADERS:
        try:
            response = requests.get(
                url, 
                headers=headers, 
                timeout=timeout,
                proxies=proxies,
                allow_redirects=True,
                verify=True
            )
            if response.status_code == 200 and 'text/html' in response.headers.get('Content-Type', ''):
                return response.text
        except requests.RequestException:
            continue
    return None


def fetch_with_cloudscraper(url, timeout):
    """Fetch with cloudscraper for Cloudflare bypass"""
    if not cloudscraper:
        return None
    try:
        proxy = get_random_proxy()
        proxies = {'http': proxy, 'https': proxy} if proxy else None
        
        scraper = cloudscraper.create_scraper(
            browser={
                'browser': 'chrome',
                'platform': 'windows',
                'mobile': False,
                'desktop': True
            },
            delay=10  # Add random delay to appear more human-like
        )
        response = scraper.get(
            url, 
            timeout=timeout,
            proxies=proxies,
            allow_redirects=True
        )
        if response.status_code == 200 and 'text/html' in response.headers.get('Content-Type', ''):
            return response.text
    except Exception as e:
        logging.warning(f"Cloudscraper fetch failed for {url}: {e}")
        return None
    return None


def fetch_with_playwright(url, timeout):
    """Enhanced Playwright fetch with anti-detection and captcha handling"""
    if not ENABLE_PLAYWRIGHT_CRAWL or not sync_playwright:
        return None
    try:
        with sync_playwright() as p:
            browser = None
            try:
                proxy = get_random_proxy()
                proxy_config = None
                if proxy:
                    # Parse proxy URL
                    from urllib.parse import urlparse
                    parsed_proxy = urlparse(proxy)
                    proxy_config = {
                        'server': f"{parsed_proxy.scheme}://{parsed_proxy.hostname}:{parsed_proxy.port}"
                    }
                    if parsed_proxy.username and parsed_proxy.password:
                        proxy_config['username'] = parsed_proxy.username
                        proxy_config['password'] = parsed_proxy.password
                
                # Launch browser with anti-detection args
                launch_options = {
                    'headless': True,
                    'args': [
                        '--disable-blink-features=AutomationControlled',
                        '--disable-dev-shm-usage',
                        '--disable-accelerated-2d-canvas',
                        '--no-first-run',
                        '--no-zygote',
                        '--disable-gpu',
                        '--no-sandbox',
                        '--disable-setuid-sandbox',
                        '--window-size=1920,1080',
                    ]
                }
                
                if proxy_config:
                    launch_options['proxy'] = proxy_config
                
                browser = p.chromium.launch(**launch_options)
                
                # Create context with realistic settings
                context = browser.new_context(
                    user_agent=HEADLESS_HEADERS[0]['User-Agent'],
                    viewport={'width': 1920, 'height': 1080},
                    locale='en-US',
                    timezone_id='America/New_York',
                    geolocation={'latitude': 40.7128, 'longitude': -74.0060},
                    permissions=['geolocation'],
                    color_scheme='light',
                    device_scale_factor=1,
                    has_touch=False,
                    is_mobile=False,
                    java_script_enabled=True,
                    extra_http_headers={
                        'Accept-Language': 'en-US,en;q=0.9',
                        'Accept-Encoding': 'gzip, deflate, br',
                        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8',
                        'Sec-Fetch-Dest': 'document',
                        'Sec-Fetch-Mode': 'navigate',
                        'Sec-Fetch-Site': 'none',
                        'Sec-Fetch-User': '?1',
                        'Upgrade-Insecure-Requests': '1'
                    }
                )
                
                page = context.new_page()
                
                # Apply stealth techniques if available
                if stealth_sync:
                    stealth_sync(page)
                
                # Add anti-detection scripts
                page.add_init_script("""
                    // Override the navigator.webdriver property
                    Object.defineProperty(navigator, 'webdriver', {
                        get: () => undefined
                    });
                    
                    // Override the navigator.plugins property
                    Object.defineProperty(navigator, 'plugins', {
                        get: () => [1, 2, 3, 4, 5]
                    });
                    
                    // Override the navigator.languages property
                    Object.defineProperty(navigator, 'languages', {
                        get: () => ['en-US', 'en']
                    });
                    
                    // Override chrome property
                    window.chrome = {
                        runtime: {}
                    };
                    
                    // Override permissions
                    const originalQuery = window.navigator.permissions.query;
                    window.navigator.permissions.query = (parameters) => (
                        parameters.name === 'notifications' ?
                            Promise.resolve({ state: Notification.permission }) :
                            originalQuery(parameters)
                    );
                """)
                
                # Navigate with retries
                max_retries = 3
                for attempt in range(max_retries):
                    try:
                        page.goto(url, wait_until='domcontentloaded', timeout=timeout * 1000)
                        break
                    except Exception as e:
                        if attempt == max_retries - 1:
                            raise e
                        import time
                        time.sleep(2)
                
                # Check for captchas
                captcha_detected = False
                captcha_type = None
                
                # Detect reCAPTCHA v2
                if page.locator('iframe[src*="recaptcha"]').count() > 0:
                    captcha_detected = True
                    captcha_type = 'recaptcha_v2'
                    logging.info(f"reCAPTCHA v2 detected on {url}")
                
                # Detect reCAPTCHA v3
                elif page.locator('script[src*="recaptcha/releases"]').count() > 0:
                    captcha_detected = True
                    captcha_type = 'recaptcha_v3'
                    logging.info(f"reCAPTCHA v3 detected on {url}")
                
                # Detect hCaptcha
                elif page.locator('iframe[src*="hcaptcha"]').count() > 0:
                    captcha_detected = True
                    captcha_type = 'hcaptcha'
                    logging.info(f"hCaptcha detected on {url}")
                
                # Attempt to bypass captcha using custom implementation
                if captcha_detected:
                    try:
                        captcha_info = {
                            'detected': True,
                            'type': captcha_type,
                            'site_key': None
                        }
                        
                        logging.info(f"Attempting to bypass {captcha_type} using custom solver...")
                        bypass_success = CustomCaptchaBypass.auto_solve_captcha(page, captcha_info, timeout=60000)
                        
                        if bypass_success:
                            logging.info(f"{captcha_type} bypassed successfully")
                        else:
                            logging.warning(f"{captcha_type} bypass failed - continuing anyway")
                    
                    except Exception as captcha_error:
                        logging.error(f"Captcha bypass error: {captcha_error}")
                
                # Scroll to load lazy-loaded images (for property listings, products, etc.)
                try:
                    # Scroll down in increments to trigger lazy loading
                    for scroll_step in range(3):
                        page.evaluate(f"window.scrollTo(0, document.body.scrollHeight * {(scroll_step + 1) / 3})")
                    # Apply stealth mode before navigation
                    CustomCaptchaBypass.apply_stealth_mode(page)
                    
                    # Navigate to the page with extended timeout
                    response = page.goto(url, wait_until='domcontentloaded', timeout=30000)
                    
                    if not response or not response.ok:
                        logging.warning(f"Failed to load {url}: HTTP {response.status if response else 'No response'}")
                        return None
                    
                    # Check for Cloudflare challenge
                    if CustomCaptchaBypass.detect_cloudflare(page):
                        logging.info("Cloudflare challenge detected - attempting bypass...")
                        cloudflare_bypassed = CustomCaptchaBypass.bypass_cloudflare(page, timeout=30000)
                        
                        if cloudflare_bypassed:
                            logging.info("Cloudflare bypass successful")
                        else:
                            logging.warning("Cloudflare bypass timeout - continuing anyway")
                    
                    # Wait for page to be fully loaded
                    page.wait_for_load_state('networkidle', timeout=15000)
                except Exception as scroll_error:
                    logging.debug(f"Scrolling warning (non-critical): {scroll_error}")
                
                # Wait for final content
                page.wait_for_timeout(2000)  # Wait 2 seconds for any final JS
                content = page.content()
                
                context.close()
                return content
            finally:
                if browser:
                    browser.close()
    except Exception as e:
        logging.warning(f"Playwright fetch failed for {url}: {e}")
        return None


def fetch_page_html(url, timeout=15):
    logging.info(f"Attempting to fetch: {url}")
    
    html = fetch_with_requests(url, timeout)
    if html:
        logging.info(f"✓ fetch_with_requests succeeded for {url}")
        return html
    else:
        logging.warning(f"✗ fetch_with_requests failed for {url}")

    html = fetch_with_cloudscraper(url, timeout)
    if html:
        logging.info(f"✓ fetch_with_cloudscraper succeeded for {url}")
        return html
    else:
        logging.warning(f"✗ fetch_with_cloudscraper failed for {url}")

    if ENABLE_PLAYWRIGHT_CRAWL:
        html = fetch_with_playwright(url, timeout)
        if html:
            logging.info(f"✓ fetch_with_playwright succeeded for {url}")
            return html
        else:
            logging.warning(f"✗ fetch_with_playwright failed for {url}")
    else:
        logging.warning(f"✗ Playwright crawling disabled (ENABLE_PLAYWRIGHT_CRAWL=false)")

    logging.error(f"✗✗✗ ALL FETCH METHODS FAILED for {url}")
    return None


def extract_text_from_html(html_content, base_url=None):
    """Convert raw HTML into clean text, extract title and images."""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Extract images before removing tags
    images = []
    for img in soup.find_all('img'):
        # Try multiple image sources (lazy loading, data attributes, srcset)
        img_url = (
            img.get('src') or 
            img.get('data-src') or 
            img.get('data-lazy-src') or 
            img.get('data-original') or
            img.get('data-lazy')
        )
        
        # Handle srcset for high-quality images
        if not img_url or 'data:image' in img_url:
            srcset = img.get('srcset', '')
            if srcset:
                # Get the largest image from srcset
                srcset_urls = [s.strip().split()[0] for s in srcset.split(',') if s.strip()]
                if srcset_urls:
                    img_url = srcset_urls[-1]  # Usually the largest
        
        if img_url and 'data:image' not in img_url:
            # Convert relative URLs to absolute
            if base_url and not img_url.startswith(('http://', 'https://')):
                img_url = urljoin(base_url, img_url)
            
            # Get alt text and class for context
            alt_text = img.get('alt', '').strip()
            img_class = img.get('class', [])
            img_class_str = ' '.join(img_class) if isinstance(img_class, list) else str(img_class)
            
            # Filter out tiny images, icons, and tracking pixels
            width = img.get('width', '')
            height = img.get('height', '')
            
            # Skip very small images (likely icons or tracking pixels)
            try:
                if width and height:
                    w = int(width.replace('px', '')) if isinstance(width, str) else int(width)
                    h = int(height.replace('px', '')) if isinstance(height, str) else int(height)
                    if w < 100 or h < 100:
                        continue
            except (ValueError, TypeError):
                pass
            
            # Skip common tracking/icon/UI patterns
            skip_patterns = [
                'tracking', 'pixel', 'icon', 'logo.svg', 'favicon', 
                'logo.png', 'logo.jpg', 'spacer.gif', '1x1',
                'navbar', 'footer', 'header-', 'banner-ad',
                'social-', 'share-button'
            ]
            if any(pattern in img_url.lower() or pattern in img_class_str.lower() for pattern in skip_patterns):
                continue
            
            # Prioritize property/product images based on common patterns
            priority = 0
            if any(x in img_class_str.lower() for x in ['property', 'listing', 'card', 'product', 'item', 'photo', 'gallery']):
                priority = 1
            
            images.append({
                'url': img_url,
                'alt': alt_text,
                'priority': priority
            })
    
    # Sort by priority (property images first) and limit
    images.sort(key=lambda x: x.get('priority', 0), reverse=True)
    
    # Remove unwanted tags
    for tag in soup(['script', 'style', 'noscript', 'iframe', 'header', 'footer', 'nav']):
        tag.decompose()

    title = soup.title.string.strip() if soup.title and soup.title.string else ""
    text_lines = [line.strip() for line in soup.get_text(separator='\n').splitlines()]
    text = "\n".join([line for line in text_lines if line])
    
    return title, text, images


def normalize_url(url):
    """Remove URL fragments and trailing slashes for consistency."""
    clean_url, _ = urldefrag(url)
    parsed = urlparse(clean_url)
    normalized = parsed._replace(fragment='', query=parsed.query).geturl()
    if normalized.endswith('/'):
        normalized = normalized[:-1]
    return normalized


def crawl_website(start_url, max_pages=5, max_depth=1, timeout=15):
    """Breadth-first crawl limited pages within the same domain."""
    parsed_start = urlparse(start_url)
    if parsed_start.scheme not in ('http', 'https'):
        raise ValueError("URL must start with http:// or https://")

    base_domain = parsed_start.netloc
    visited = set()
    queue = deque([(normalize_url(start_url), 0)])
    pages = []
    while queue and len(pages) < max_pages:
        current_url, depth = queue.popleft()
        if current_url in visited or depth > max_depth:
            continue
        visited.add(current_url)

        try:
            html = fetch_page_html(current_url, timeout=timeout)
            if not html:
                continue
            title, text, images = extract_text_from_html(html, current_url)
            if not text.strip():
                continue

            pages.append({
                'url': current_url,
                'title': title or current_url,
                'text': text,
                'images': images,  # Include extracted images
                'depth': depth
            })

            if depth >= max_depth:
                continue

            soup = BeautifulSoup(html, 'html.parser')
            for anchor in soup.find_all('a', href=True):
                linked_url = urljoin(current_url, anchor['href'])
                normalized_link = normalize_url(linked_url)
                parsed_link = urlparse(normalized_link)
                if parsed_link.netloc != base_domain:
                    continue
                if normalized_link in visited:
                    continue
                queue.append((normalized_link, depth + 1))
                if len(queue) + len(pages) >= max_pages:
                    break
        except requests.RequestException:
            continue

    return pages


@app.route('/', methods=['GET'])
def root():
    """Root endpoint"""
    return jsonify({
        'message': 'Knowledge Base API is running!',
        'status': 'online',
        'endpoints': {
            'health': '/health',
            'upload': '/api/upload',
            'documents': '/api/documents',
            'indexes': '/api/indexes'
        }
    })


@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'pinecone_configured': bool(PINECONE_API_KEY),
        'gemini_configured': bool(GEMINI_API_KEY),
        'embedding_model': 'Google Gemini text-embedding-004',
        'embedding_type': 'API - Matches n8n workflow',
        'embedding_dimension': EMBEDDING_DIMENSION,
        'quota_limits': 'Google Gemini API limits apply'
    })


@app.route('/api/upload', methods=['POST'])
def upload_document():
    """Upload and index document to Pinecone"""
    try:
        # Validate request
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        project = request.form.get('project', 'default')
        index_name = request.form.get('index_name', DEFAULT_INDEX_NAME)
        
        if file.filename == '':
            return jsonify({'error': 'Empty filename'}), 400
        
        # Read file
        file_bytes = file.read()
        filename = file.filename
        
        # Extract text
        text = extract_text(file_bytes, filename)
        
        if not text or len(text.strip()) == 0:
            return jsonify({'error': 'No text extracted from file'}), 400
        
        ingest_result = ingest_text_payload(
            text,
            filename,
            project,
            index_name,
            extra_metadata={
                'file_size': len(file_bytes),
                'source': filename,
                'content_type': 'file_upload'
            }
        )
        
        return jsonify({
            'success': True,
            **ingest_result
        })
        
    except Exception as e:
        print(f"Upload error: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/ingest-url', methods=['POST'])
def ingest_url():
    """Crawl a URL (and optional child pages) and ingest the content."""
    try:
        data = request.get_json(force=True) or {}
        start_url = (data.get('url') or '').strip()
        project = data.get('project', 'default')
        index_name = data.get('index_name', DEFAULT_INDEX_NAME)
        max_pages = data.get('max_pages', 5)
        max_depth = data.get('max_depth', 1)

        if not start_url:
            return jsonify({'error': 'URL is required'}), 400

        try:
            max_pages = max(1, min(int(max_pages), 25))
        except (TypeError, ValueError):
            max_pages = 5

        try:
            max_depth = max(0, min(int(max_depth), 3))
        except (TypeError, ValueError):
            max_depth = 1

        pages = crawl_website(start_url, max_pages=max_pages, max_depth=max_depth)

        if not pages:
            return jsonify({'error': 'No crawlable pages were found at the supplied URL.'}), 400

        ingested = []
        skipped = 0

        for idx, page in enumerate(pages, start=1):
            try:
                # Prepare image URLs for metadata
                image_urls = [img['url'] for img in page.get('images', [])[:10]]  # Limit to 10 images per page
                image_alts = [img['alt'] for img in page.get('images', [])[:10]]
                
                # Add image context to text for better search
                text_with_images = page['text']
                if image_urls:
                    image_context = "\n\n[Page Images]: " + ", ".join([f"{alt or 'Image'}" for alt in image_alts if alt])
                    text_with_images += image_context
                
                ingest_result = ingest_text_payload(
                    text_with_images,
                    page['title'],
                    project,
                    index_name,
                    extra_metadata={
                        'file_size': len(page['text']),
                        'source': page['title'],
                        'source_url': page['url'],
                        'crawl_depth': page['depth'],
                        'content_type': 'web_page',
                        'page_index': idx,
                        'image_urls': json.dumps(image_urls) if image_urls else '',  # Store as JSON string
                        'image_count': len(image_urls)
                    }
                )
                ingested.append({
                    'url': page['url'],
                    'title': page['title'],
                    'depth': page['depth'],
                    'chunks_created': ingest_result['chunks_created'],
                    'total_characters': ingest_result['total_characters'],
                    'images_found': len(image_urls)
                })
            except Exception as page_error:
                print(f"URL ingest error for {page['url']}: {page_error}")
                skipped += 1
                continue

        if not ingested:
            return jsonify({'error': 'Failed to ingest any pages from the supplied URL.'}), 500

        return jsonify({
            'success': True,
            'pages_ingested': len(ingested),
            'pages_skipped': skipped,
            'total_chunks': sum(page['chunks_created'] for page in ingested),
            'details': ingested
        })

    except Exception as e:
        print(f"URL ingest error: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/search', methods=['POST'])
def search_documents():
    """Search documents using semantic search"""
    try:
        data = request.get_json()
        query = data.get('query', '')
        project = data.get('project', 'default')
        index_name = data.get('index_name', DEFAULT_INDEX_NAME)
        top_k = data.get('top_k', 10)
        
        if not query:
            return jsonify({'error': 'Query is required'}), 400
        
        # Generate query embedding
        query_embedding = generate_embeddings([query])[0]
        
        # Search in Pinecone
        index = get_or_create_index(index_name)
        results = index.query(
            vector=query_embedding,
            top_k=top_k,
            namespace=project,
            include_metadata=True
        )
        
        # Format results
        search_results = []
        for match in results.matches:
            search_results.append({
                'id': match.id,
                'score': float(match.score),
                'metadata': match.metadata
            })
        
        return jsonify({
            'success': True,
            'query': query,
            'results': search_results,
            'total_results': len(search_results)
        })
        
    except Exception as e:
        print(f"Search error: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/chat', methods=['POST'])
def chat_with_knowledge_base():
    """Chat with the knowledge base using RAG"""
    try:
        if not GEMINI_API_KEY:
            return jsonify({'error': 'GEMINI_API_KEY is not configured on the server'}), 500

        data = request.get_json(force=True) or {}
        query = data.get('query', '').strip()
        project = data.get('project', 'default')
        index_name = data.get('index_name', DEFAULT_INDEX_NAME)
        top_k = data.get('top_k', 3)
        history = data.get('history', [])

        if not query:
            return jsonify({'error': 'Query is required'}), 400

        try:
            top_k = int(top_k)
        except (TypeError, ValueError):
            top_k = 3
        top_k = max(1, min(top_k, 20))

        # Lazily initialize chat model if needed (handles hot reloads)
        global chat_model, active_chat_model_name
        if chat_model is None:
            chat_model = initialize_chat_model()

        # Embed query and retrieve context from Pinecone
        query_embedding = generate_embeddings([query])[0]
        index = get_or_create_index(index_name)
        results = index.query(
            vector=query_embedding,
            top_k=top_k,
            namespace=project,
            include_metadata=True
        )

        if not results.matches:
            return jsonify({
                'success': True,
                'answer': "I couldn't find any relevant information in the selected repo. Please try rephrasing your question or choose another repo.",
                'sources': []
            })

        max_context_chars = int(data.get('max_context_chars', 6000))
        context_sections = []
        sources = []
        running_chars = 0

        for idx, match in enumerate(results.matches, start=1):
            metadata = match.metadata or {}
            passage = (metadata.get('text') or '').strip()
            if not passage:
                continue

            remaining_budget = max_context_chars - running_chars
            if remaining_budget <= 0:
                break

            if len(passage) > remaining_budget:
                passage = passage[:remaining_budget] + '...'

            # Extract image URLs if available
            image_urls_json = metadata.get('image_urls', '')
            image_urls = []
            if image_urls_json:
                try:
                    image_urls = json.loads(image_urls_json)
                except (json.JSONDecodeError, TypeError):
                    pass
            
            source_url = metadata.get('source_url', '')
            section_header = f"[Source {idx}] {metadata.get('filename', 'Unknown file')} (chunk {metadata.get('chunk_index', '-')})"
            
            # Add images to context if available
            context_content = f"{section_header}\n{passage}"
            if image_urls and source_url:
                context_content += f"\n[Images from {source_url}]: {', '.join(image_urls[:3])}"  # Include up to 3 image URLs
            
            context_sections.append(context_content)

            sources.append({
                'source_id': idx,
                'score': float(match.score),
                'document_id': metadata.get('document_id'),
                'filename': metadata.get('filename'),
                'chunk_index': metadata.get('chunk_index'),
                'text': passage,
                'source_url': source_url,
                'image_urls': image_urls[:3] if image_urls else []  # Include up to 3 images in response
            })

            running_chars += len(passage)

        if not context_sections:
            return jsonify({
                'success': True,
                'answer': "I couldn't find any relevant passages in the selected repo for that question.",
                'sources': []
            })

        formatted_history = []
        for turn in history[-6:]:
            user_turn = turn.get('user', '').strip()
            assistant_turn = turn.get('assistant', '').strip()
            if user_turn:
                formatted_history.append(f"User: {user_turn}")
            if assistant_turn:
                formatted_history.append(f"Assistant: {assistant_turn}")
        conversation_context = "\n".join(formatted_history) if formatted_history else "No previous conversation."
        context_block = "\n\n".join(context_sections)

        prompt = (
            "You are Habib, a friendly property consultant helping users find properties in Singapore.\n"
            "Use the knowledge base context to provide accurate property information.\n\n"
            "FORMATTING RULES:\n"
            "- Use clean, natural language without excessive formatting symbols\n"
            "- Number properties as: 1. 2. 3.\n"
            "- For each property, use this structure:\n"
            "  Property Name at Location\n"
            "  Price: [amount]\n"
            "  Type: [property type]\n"
            "  Size: [sqft], Bedrooms: [number], Bathrooms: [number]\n"
            "  Location: [full address]\n"
            "  MRT: [distance and station]\n"
            "  Built: [year]\n\n"
            "IMPORTANT:\n"
            "- DO NOT use asterisks, bold, or markdown formatting\n"
            "- DO NOT include image links - images are displayed separately\n"
            "- Provide 2-3 property options when asked for suggestions\n"
            "- Only use information from the knowledge base - never invent details\n"
            "- Keep responses conversational and end with a helpful follow-up question\n\n"
            f"Conversation history:\n{conversation_context}\n\n"
            f"Knowledge base:\n{context_block}\n\n"
            f"User: {query}\n\n"
            "Provide a helpful response with property details in clean, readable format."
        )

        try:
            response = chat_model.generate_content(prompt)
        except NotFound:
            # Fallback once more in case the remote model registry changed between requests
            chat_model = initialize_chat_model(force_fallback=True)
            response = chat_model.generate_content(prompt)
        except GoogleAPICallError as api_error:
            return jsonify({'error': f'Gemini API error: {api_error.message}'}), 500
        answer_text = getattr(response, 'text', None)
        if not answer_text and hasattr(response, 'candidates'):
            for candidate in response.candidates:
                if hasattr(candidate, 'content') and candidate.content:
                    parts = getattr(candidate.content, 'parts', None)
                    if parts:
                        answer_text = " ".join(getattr(part, 'text', '') for part in parts if getattr(part, 'text', ''))
                        if answer_text:
                            break

        if not answer_text:
            answer_text = "I encountered an issue while generating a response. Please try again."

        # Collect all unique images from sources for easy frontend rendering
        all_images = []
        seen_urls = set()
        for source in sources:
            for img_url in source.get('image_urls', []):
                if img_url and img_url not in seen_urls and 'logo' not in img_url.lower():
                    seen_urls.add(img_url)
                    all_images.append({
                        'url': img_url,
                        'source_url': source.get('source_url', ''),
                        'title': source.get('filename', '')
                    })
        
        return jsonify({
            'success': True,
            'answer': answer_text.strip(),
            'sources': sources,
            'property_images': all_images[:6]  # Limit to 6 images for display
        })

    except Exception as e:
        print(f"Chat error: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents', methods=['GET'])
def list_documents():
    """List all documents in a project"""
    try:
        project = request.args.get('project', 'default')
        index_name = request.args.get('index_name', DEFAULT_INDEX_NAME)
        
        # Get index stats
        index = get_or_create_index(index_name)
        index_dimension = get_index_dimension(index_name)
        stats = index.describe_index_stats()
        
        # Get namespace stats
        namespace_stats = stats.namespaces.get(project, {})
        
        # Query for sample documents to get unique document IDs
        # This is a workaround since Pinecone doesn't have a direct "list all docs" API
        sample_results = index.query(
            vector=[0.0] * index_dimension,
            top_k=10000,
            namespace=project,
            include_metadata=True
        )
        
        # Group by document_id
        documents = {}
        for match in sample_results.matches:
            doc_id = match.metadata.get('document_id')
            if doc_id not in documents:
                documents[doc_id] = {
                    'document_id': doc_id,
                    'filename': match.metadata.get('filename'),
                    'project': match.metadata.get('project'),
                    'upload_date': match.metadata.get('upload_date'),
                    'total_chunks': match.metadata.get('total_chunks', 0),
                    'file_size': match.metadata.get('file_size', 0)
                }
        
        documents_list = list(documents.values())
        documents_list.sort(key=lambda x: x.get('upload_date', ''), reverse=True)
        
        return jsonify({
            'success': True,
            'project': project,
            'documents': documents_list,
            'total_documents': len(documents_list),
            'total_vectors': namespace_stats.get('vector_count', 0)
        })
        
    except Exception as e:
        print(f"List documents error: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/<document_id>', methods=['DELETE'])
def delete_document(document_id):
    """Delete a document and all its chunks"""
    try:
        project = request.args.get('project', 'default')
        index_name = request.args.get('index_name', DEFAULT_INDEX_NAME)
        
        index = get_or_create_index(index_name)
        index_dimension = get_index_dimension(index_name)
        
        # Find all chunks for this document
        results = index.query(
            vector=[0.0] * index_dimension,
            top_k=10000,
            namespace=project,
            include_metadata=True,
            filter={'document_id': document_id}
        )
        
        # Delete all chunks
        vector_ids = [match.id for match in results.matches]
        if vector_ids:
            index.delete(ids=vector_ids, namespace=project)
        
        return jsonify({
            'success': True,
            'document_id': document_id,
            'chunks_deleted': len(vector_ids)
        })
        
    except Exception as e:
        print(f"Delete error: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/stats', methods=['GET'])
def get_stats():
    """Get overall statistics"""
    try:
        index_name = request.args.get('index_name', DEFAULT_INDEX_NAME)
        index = get_or_create_index(index_name)
        stats = index.describe_index_stats()
        
        return jsonify({
            'success': True,
            'total_vectors': stats.total_vector_count,
            'namespaces': {
                name: {
                    'vector_count': ns.vector_count
                }
                for name, ns in stats.namespaces.items()
            }
        })
        
    except Exception as e:
        print(f"Stats error: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/indexes', methods=['GET'])
def list_indexes():
    """List all Pinecone indexes"""
    try:
        indexes = pc.list_indexes()
        index_list = []
        
        for index_info in indexes:
            index_list.append({
                'name': index_info.name,
                'dimension': index_info.dimension,
                'metric': index_info.metric,
                'host': index_info.host
            })
        
        return jsonify({
            'success': True,
            'indexes': index_list,
            'total_indexes': len(index_list)
        })
        
    except Exception as e:
        print(f"List indexes error: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/indexes', methods=['POST'])
def create_index():
    """Create a new Pinecone index"""
    try:
        data = request.get_json()
        index_name = data.get('index_name', '').strip()
        
        if not index_name:
            return jsonify({'error': 'Index name is required'}), 400
        
        # Validate index name (alphanumeric and hyphens only)
        if not index_name.replace('-', '').replace('_', '').isalnum():
            return jsonify({'error': 'Index name must contain only alphanumeric characters, hyphens, and underscores'}), 400
        
        # Check if index already exists
        existing_indexes = [idx.name for idx in pc.list_indexes()]
        
        if index_name in existing_indexes:
            return jsonify({'error': f'Index "{index_name}" already exists'}), 400
        
        # Create the index
        pc.create_index(
            name=index_name,
            dimension=EMBEDDING_DIMENSION,
            metric='cosine',
            spec=ServerlessSpec(
                cloud='aws',
                region=PINECONE_ENVIRONMENT
            )
        )
        
        return jsonify({
            'success': True,
            'message': f'Index "{index_name}" created successfully',
            'index_name': index_name,
            'dimension': EMBEDDING_DIMENSION
        })
        
    except Exception as e:
        print(f"Create index error: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/indexes/<index_name>', methods=['DELETE'])
def delete_index(index_name):
    """Delete a Pinecone index"""
    try:
        # Prevent deletion of default index
        if index_name == DEFAULT_INDEX_NAME:
            return jsonify({'error': f'Cannot delete default index "{DEFAULT_INDEX_NAME}"'}), 400
        
        # Check if index exists
        existing_indexes = [idx.name for idx in pc.list_indexes()]
        
        if index_name not in existing_indexes:
            return jsonify({'error': f'Index "{index_name}" does not exist'}), 404
        
        # Delete the index
        pc.delete_index(index_name)
        
        return jsonify({
            'success': True,
            'message': f'Index "{index_name}" deleted successfully',
            'index_name': index_name
        })
        
    except Exception as e:
        print(f"Delete index error: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/analyse-document', methods=['POST'])
def analyse_document():
    """Analyze documents (Excel, PDF, DOCX, TXT) with AI"""
    try:
        if not GEMINI_API_KEY:
            return jsonify({'error': 'GEMINI_API_KEY is not configured on the server'}), 500
        
        file_type = request.form.get('file_type', '').lower()
        file_type = request.form.get('file_type')
        conversation_history_json = request.form.get('conversation_history', '[]')
        
        try:
            conversation_history = json.loads(conversation_history_json)
        except:
            conversation_history = []
        
        if file_type == 'url':
            # Handle URL analysis
            url = request.form.get('url')
            if not url:
                return jsonify({'error': 'URL is required'}), 400
            
            try:
                from bs4 import BeautifulSoup
                
                # Try cloudscraper first (bypasses Cloudflare and bot protection)
                try:
                    import cloudscraper
                    scraper = cloudscraper.create_scraper(
                        browser={
                            'browser': 'chrome',
                            'platform': 'windows',
                            'mobile': False
                        }
                    )
                    response = scraper.get(url, timeout=15)
                    response.raise_for_status()
                except:
                    # Fallback to requests with comprehensive headers
                    import requests
                    headers_req = {
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                        'Accept-Language': 'en-US,en;q=0.9',
                        'Accept-Encoding': 'gzip, deflate, br',
                        'DNT': '1',
                        'Connection': 'keep-alive',
                        'Upgrade-Insecure-Requests': '1'
                    }
                    response = requests.get(url, headers=headers_req, timeout=15, allow_redirects=True)
                    response.raise_for_status()
                
                # Parse HTML
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # Remove unwanted elements
                for element in soup(['script', 'style', 'nav', 'header', 'footer', 'aside']):
                    element.decompose()
                
                # Extract text
                text = soup.get_text()
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text_content = ' '.join(chunk for chunk in chunks if chunk)[:5000]
                
                # Check if we got meaningful content
                if len(text_content.strip()) < 100:
                    return jsonify({'error': 'Unable to extract content from this webpage. It may be protected or require JavaScript.'}), 400
                
                prompt = f"""Analyze this webpage from {url}:

{text_content}

CRITICAL: Maximum 50 words.

• [Key point 1]
• [Key point 2]
• [Key point 3]"""
                
                chat_model = initialize_chat_model()
                ai_response = chat_model.generate_content(prompt)
                analysis = ai_response.text if hasattr(ai_response, 'text') else str(ai_response)
                
                return jsonify({'success': True, 'analysis': analysis, 'text_content': text_content})
            except Exception as e:
                error_msg = str(e)
                if '403' in error_msg or 'Forbidden' in error_msg:
                    return jsonify({'error': 'This website blocks automated access. Try a different URL or use the file upload feature.'}), 403
                elif '404' in error_msg:
                    return jsonify({'error': 'Page not found. Please check the URL.'}), 404
                elif 'timeout' in error_msg.lower():
                    return jsonify({'error': 'Request timeout. The website took too long to respond.'}), 408
                else:
                    return jsonify({'error': f'Unable to fetch webpage: {error_msg}'}), 400
        
        # Handle follow-up questions
        question = request.form.get('question')
        if question:
            if file_type == 'excel':
                excel_data_json = request.form.get('excel_data', '[]')
                try:
                    excel_data = json.loads(excel_data_json)
                except:
                    return jsonify({'error': 'Invalid Excel data format'}), 400
                
                headers = excel_data[0] if len(excel_data) > 0 else []
                sample_rows = excel_data[1:6] if len(excel_data) > 1 else []
                
                context = f"""You are a data analyst. Excel data:
Columns: {', '.join(str(h) for h in headers)}
Rows: {len(excel_data) - 1}

Sample:"""
                for i, row in enumerate(sample_rows, 1):
                    context += f"\nRow {i}: {dict(zip(headers, row))}"
                
                if conversation_history:
                    context += "\n\nConversation:\n"
                    for turn in conversation_history[-3:]:
                        context += f"User: {turn.get('user', '')}\nAssistant: {turn.get('assistant', '')}\n"
                
                context += f"\n\nQuestion: {question}\nAnswer in max 20 words."
                
            elif file_type == 'text':
                text_content = request.form.get('text_content', '')
                context = f"""Text: {text_content[:2000]}..."""
                if conversation_history:
                    context += "\nConversation:\n"
                    for turn in conversation_history[-3:]:
                        context += f"User: {turn.get('user', '')}\nAI: {turn.get('assistant', '')}\n"
                context += f"\nQuestion: {question}\nAnswer in max 20 words."
            
            elif file_type == 'url':
                text_content = request.form.get('text_content', '')
                context = f"""Webpage content: {text_content[:2000]}..."""
                if conversation_history:
                    context += "\nConversation:\n"
                    for turn in conversation_history[-3:]:
                        context += f"User: {turn.get('user', '')}\nAI: {turn.get('assistant', '')}\n"
                context += f"\nQuestion: {question}\nAnswer in max 20 words."
            
            else:
                context = f"Question: {question}\nAnswer briefly."
            
            chat_model = initialize_chat_model()
            response = chat_model.generate_content(context)
            answer = response.text if hasattr(response, 'text') else str(response)
            
            return jsonify({'success': True, 'answer': answer})
        
        # Initial analysis (no question provided)
        file = request.files.get('file')
        if not file:
            return jsonify({'error': 'No file provided'}), 400
        
        if file_type == 'excel':
            excel_data_json = request.form.get('excel_data', '{}')
            try:
                excel_metadata = json.loads(excel_data_json)
            except:
                return jsonify({'error': 'Invalid Excel metadata'}), 400
            
            rows = excel_metadata.get('rows', 0)
            columns = excel_metadata.get('columns', 0)
            headers = excel_metadata.get('headers', [])
            sample_data = excel_metadata.get('sample_data', [])
            
            # Create analysis prompt
            prompt = f"""You are an expert data analyst. Analyze this Excel spreadsheet:

Rows: {rows}
Columns: {columns}
Column headers: {', '.join(str(h) for h in headers)}

Sample data (first 5 rows):
"""
            for i, row in enumerate(sample_data, 1):
                prompt += f"\nRow {i}: {dict(zip(headers, row))}"
            
            prompt += """

CRITICAL: BE EXTREMELY BRIEF. Maximum 50 words total.

Format (EXACTLY):
"This is [type] data with [X] rows.

• [One key finding with number]
• [Second key finding with number]
• [Third key finding with number]

SUGGESTED CHARTS:
[Chart Type 1]: [Why - 5 words max]
[Chart Type 2]: [Why - 5 words max]"

RULES:
- NO paragraphs, NO explanations, NO calculations
- Just 3 bullet points (each max 8 words)
- Then suggest 1-2 appropriate chart types
- Chart types: Pie Chart, Line Chart, Bar Chart, or Scatter Plot
- Briefly explain why each chart fits"""
            
            # Generate analysis
            chat_model = initialize_chat_model()
            response = chat_model.generate_content(prompt)
            analysis = response.text if hasattr(response, 'text') else str(response)
            
            # Infer column types
            column_types = {}
            if len(sample_data) > 0:
                for i, header in enumerate(headers):
                    sample_values = [row[i] for row in sample_data if i < len(row)]
                    if all(isinstance(v, (int, float)) for v in sample_values if v):
                        column_types[str(header)] = 'numeric'
                    else:
                        column_types[str(header)] = 'text'
            
            return jsonify({
                'success': True,
                'analysis': analysis,
                'data_summary': {
                    'rows': rows,
                    'columns': columns,
                    'column_types': column_types
                }
            })
        
        elif file_type == 'pdf':
            # Extract text from PDF
            try:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ''
                for page in pdf_reader.pages[:10]:  # First 10 pages
                    text_content += page.extract_text()
                
                prompt = f"""You are a document analyst. Analyze this PDF document:

{text_content[:3000]}...

IMPORTANT: Keep your analysis BRIEF and ACTIONABLE.

Provide ONLY:
1. One sentence: what is this document?
2. Top 2-3 KEY POINTS ONLY (bullet points)

Format:
"This is a [type] document about [topic].

• [Key point 1]
• [Key point 2]
• [Key point 3]"

Keep it short - users can ask for details."""
                
                chat_model = initialize_chat_model()
                response = chat_model.generate_content(prompt)
                analysis = response.text if hasattr(response, 'text') else str(response)
                
                return jsonify({
                    'success': True,
                    'analysis': analysis
                })
            except Exception as e:
                return jsonify({'error': f'PDF parsing error: {str(e)}'}), 400
        
        elif file_type == 'text':
            text_content = request.form.get('text_content', '')
            if not text_content:
                text_content = file.read().decode('utf-8')
            
            prompt = f"""You are a document analyst. Analyze this text document:

{text_content[:3000]}...

IMPORTANT: Keep your analysis BRIEF.

Provide ONLY:
1. One sentence: what is this?
2. Top 2-3 KEY POINTS (bullet points)

Format:
"This document is about [topic].

• [Key point 1]
• [Key point 2]
• [Key point 3]"

Short and clear."""
            
            chat_model = initialize_chat_model()
            response = chat_model.generate_content(prompt)
            analysis = response.text if hasattr(response, 'text') else str(response)
            
            return jsonify({
                'success': True,
                'analysis': analysis
            })
        
        elif file_type == 'docx':
            # Extract text from DOCX
            try:
                doc = docx.Document(file)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                prompt = f"""You are a document analyst. Analyze this Word document:

{text_content[:3000]}...

IMPORTANT: Keep your analysis BRIEF.

Provide ONLY:
1. One sentence: what is this?
2. Top 2-3 KEY POINTS (bullet points)

Format:
"This is a [type] document about [topic].

• [Key point 1]
• [Key point 2]
• [Key point 3]"

Short and actionable."""
                
                chat_model = initialize_chat_model()
                response = chat_model.generate_content(prompt)
                analysis = response.text if hasattr(response, 'text') else str(response)
                
                return jsonify({
                    'success': True,
                    'analysis': analysis
                })
            except Exception as e:
                return jsonify({'error': f'DOCX parsing error: {str(e)}'}), 400
        
        else:
            return jsonify({'error': 'Unsupported file type'}), 400
    
    except Exception as e:
        print(f"Document analysis error: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/')
def health_check():
    """Health check endpoint for Railway/deployment platforms"""
    return jsonify({
        'status': 'healthy',
        'service': 'Knowledge Base API',
        'version': '1.0.0',
        'playwright_enabled': ENABLE_PLAYWRIGHT_CRAWL
    }), 200

if __name__ == '__main__':
    # Get port from environment variable (required for Render)
    port = int(os.environ.get('PORT', 5001))
    print(f"Starting server on port {port}...")
    app.run(debug=False, host='0.0.0.0', port=port)
